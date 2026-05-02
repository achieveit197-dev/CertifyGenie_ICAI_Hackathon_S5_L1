from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins (narrow for more content per page) ---
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# --- Styles helpers ---
def set_font(run, name="Calibri", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, name="Calibri", size=12 if level == 1 else 11,
             bold=True, color=(0, 51, 102))
    return p

def add_body(doc, text, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        set_font(run, size=10.5)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 90)
    set_font(run, size=7, color=(150, 150, 150))
    return p

# ═══════════════════════════════════════════
# HEADER BLOCK
# ═══════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(2)
r = p_title.add_run("CERTIFY GENIE")
set_font(r, name="Calibri", size=20, bold=True, color=(0, 51, 102))

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(2)
r2 = p_sub.add_run("AI-Powered CA Certificate Generation from Audited Financial Statements")
set_font(r2, name="Calibri", size=12, italic=True, color=(60, 60, 60))

p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_meta.paragraph_format.space_after = Pt(6)
r3 = p_meta.add_run("CA Gopi  |  ZenCFO, Chennai  |  ICAI AI Hackathon — Season 5")
set_font(r3, name="Calibri", size=10, color=(100, 100, 100))

add_divider(doc)

# ═══════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════
add_heading(doc, "1.  Introduction")
add_body(doc,
    "Issuing a Net Worth or Turnover Certificate should be a five-minute task. In practice, a CA "
    "manually reads a scanned balance sheet, types figures into a Word template, cross-checks "
    "annexures, and reformats — a process that takes 45–60 minutes and is error-prone. "
    "Certify Genie compresses that workflow to under three minutes: the CA uploads the document, "
    "Claude AI extracts the figures, the system validates the accounting equations, and a "
    "professional ICAI-format certificate is generated in PDF and DOCX. The entire pipeline runs "
    "PII redaction before the document reaches the AI, keeping the firm compliant with India's "
    "Digital Personal Data Protection Act 2023.")

# ═══════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ═══════════════════════════════════════════
add_heading(doc, "2.  Problem Statement")
add_body(doc,
    "Certificate issuance at a CA firm involves three compounding risks:")
add_bullet(doc, "Manual transcription errors — figures typed from a scanned document are frequently transposed or misread, yet no systemic check exists to flag when the final Net Worth on the certificate does not reconcile to Share Capital + Reserves − Losses.")
add_bullet(doc, "Wrong calculation method — Net Worth can be computed under Book Value, Tangible Net Worth, Companies Act Sec 2(57), SEBI LODR, NAV, or other bases. Choosing the incorrect method for the entity type or purpose creates a professional risk.")
add_bullet(doc, "PII exposure — client balance sheets contain CIN, PAN, GSTIN, and Aadhaar-linked details. Forwarding the raw document to any external service is a DPDP Act 2023 exposure without a deliberate redaction step.")

# ═══════════════════════════════════════════
# 3. SOLUTION
# ═══════════════════════════════════════════
add_heading(doc, "3.  Solution Overview")
add_body(doc,
    "Certify Genie is a browser-based five-step wizard. The CA uploads a document, reviews AI-extracted "
    "figures (with confidence scores and exact source lines), enters firm details, selects the "
    "calculation method, edits the draft certificate live on screen, and downloads the final output. "
    "The CA never types a single number from the document — the AI reads it and presents it for review.")

# ── Key Capabilities (2-col table) ──
add_heading(doc, "Key Capabilities", level=2)

table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = [("Module", "What It Does")]
rows_data = [
    ("PII Redaction Engine",
     "Blacks out CIN, PAN, GSTIN, Aadhaar, mobile, email via PyMuPDF before any content reaches Claude. CA can draw additional manual redaction boxes. Redactions are permanent, not hidden."),
    ("AI Extraction (Claude Vision)",
     "Detects text-native vs. scanned PDFs automatically. For scans, pages are rendered to PNG and sent via Claude's vision API. Returns structured JSON with 12+ financial fields, confidence scores, and source lines."),
    ("Arithmetic Validation",
     "Validates three equations — Net Worth = SC + R&S − Losses; Total Assets = External Liabilities + Equity; Working Capital = CA − CL — each with 1% tolerance. Flags discrepancies before certificate issue."),
    ("Net Worth Method Selector",
     "Computes Net Worth under all 7 methods (Book Value, Tangible NW, Companies Act Sec 2(57), SEBI LODR, NAV, Adjusted Tangible BV, Consolidated NW). Recommends the appropriate method based on entity type."),
    ("Certificate Generation",
     "Outputs ICAI-format PDF (ReportLab) and DOCX (python-docx) with letterhead, figures table, UDIN field, and signature block. Filenames are meaningful. Optional PDF password encryption."),
]

# Header row
hdr_cells = table.rows[0].cells
for i, (h, _) in enumerate(headers):
    pass

# Rebuild table with correct row count
table = doc.add_table(rows=1 + len(rows_data), cols=2)
table.style = 'Table Grid'

# Header
hdr = table.rows[0].cells
for cell, text in zip(hdr, ["Module", "Capability"]):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    set_font(run, bold=True, size=10, color=(255, 255, 255))
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '003366')
    tcPr.append(shd)

for i, (mod, cap) in enumerate(rows_data):
    row = table.rows[i + 1].cells
    row[0].text = mod
    row[1].text = cap
    for cell in row:
        for run in cell.paragraphs[0].runs:
            set_font(run, size=9.5)
    # Alternating row shading
    if i % 2 == 0:
        for cell in row:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'EEF2F7')
            tcPr.append(shd)

# Column widths
from docx.shared import Cm
for row in table.rows:
    row.cells[0].width = Cm(4.5)
    row.cells[1].width = Cm(11.5)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ═══════════════════════════════════════════
# 4. PROCESS FLOW
# ═══════════════════════════════════════════
add_heading(doc, "4.  How It Works — Five Steps")

steps = [
    ("Upload", "CA drops a PDF or Excel balance sheet. System auto-detects if the document is text-native or scanned."),
    ("Redact", "Auto-redaction runs on PII fields. CA can optionally draw manual redaction boxes before proceeding."),
    ("Extract & Validate", "Claude reads the document and returns structured financial data. The validation engine checks all three accounting equations and flags any mismatch."),
    ("Select Method", "CA selects entity type; system recommends the appropriate Net Worth method and computes the figure under all applicable bases."),
    ("Preview, Edit & Download", "Live certificate preview updates as the CA types. UDIN, notes, and PDF password can be added. PDF + DOCX are generated simultaneously with a built-in AI extraction trail."),
]

for num, (step, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    r_num = p.add_run(f"  Step {num}:  ")
    set_font(r_num, bold=True, size=10.5, color=(0, 51, 102))
    r_step = p.add_run(f"{step} — ")
    set_font(r_step, bold=True, size=10.5)
    r_desc = p.add_run(desc)
    set_font(r_desc, size=10.5)

# ═══════════════════════════════════════════
# 5. IMPACT
# ═══════════════════════════════════════════
add_heading(doc, "5.  Impact & Benefits")

impacts = [
    ("Time saved", "45–60 min → under 3 min per certificate. For a firm issuing 15–20 certificates/month, that is 8–12 hours of professional time reclaimed monthly."),
    ("Accuracy", "Arithmetic validation catches equation mismatches before issue — rounding differences, prior-year entries, or extraction errors surface immediately rather than reaching the recipient silently."),
    ("Method compliance", "Seven Net Worth methods with entity-aware defaults mean the CA is far less likely to certify on the wrong basis for a bank or tender submission."),
    ("PII discipline", "Redaction is a non-optional workflow step, not a manual reminder. The CA sees a summary of what was redacted (CIN: 1, PAN: 2…) before extraction runs."),
    ("Built-in audit trail", "The download screen shows every figure alongside the exact source line it was read from — a working paper by construction."),
]

for title, detail in impacts:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"• {title}: ")
    set_font(r1, bold=True, size=10.5)
    r2 = p.add_run(detail)
    set_font(r2, size=10.5)

# ═══════════════════════════════════════════
# 6. TECH STACK (compact inline)
# ═══════════════════════════════════════════
add_heading(doc, "6.  Technology Stack")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run(
    "Backend: Python 3.11 · FastAPI · PyMuPDF · Anthropic Claude claude-sonnet-4-6 (text + vision) · ReportLab · "
    "python-docx · pypdf · Pydantic\n"
    "Frontend: React 18 · TypeScript · Vite · Tailwind CSS · Framer Motion · Lucide React\n"
    "Compliance: PII redaction (CIN, PAN, GSTIN, Aadhaar, Mobile, Email) · No document persistence · "
    "In-memory session storage · DPDP Act 2023 data minimisation"
)
set_font(r, size=10)

# ═══════════════════════════════════════════
# 7. FUTURE SCOPE
# ═══════════════════════════════════════════
add_heading(doc, "7.  Future Scope")
future = [
    "UDIN Portal integration — direct API connection to icaiudin.org",
    "Multi-document ingestion — full financial statement set (BS, P&L, Cash Flow, Schedules)",
    "ICAI UDIN register linkage — firm-level certificate tracking for quality review",
    "Auto-generated working paper alongside the certificate",
    "Authentication & multi-user deployment — Google OAuth, PostgreSQL, per-CA history",
    "Direct submission to bank/GeM portals via API integration",
]
for item in future:
    add_bullet(doc, item)

add_divider(doc)

# ═══════════════════════════════════════════
# CONCLUSION
# ═══════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Conclusion")
set_font(r, bold=True, size=11, color=(0, 51, 102))

add_body(doc,
    "The ICAI has consistently held that technology should enhance the CA's role — removing mechanical "
    "work, not professional judgment. Certify Genie applies that principle to the certificate: the "
    "arithmetic, the extraction, the formatting, the cross-checking, and the PII handling are automated. "
    "What remains is the CA's review, method judgment, and signature — exactly where professional "
    "accountability belongs. The tool is live and generating real certificates today. "
    "The path forward — UDIN integration, working paper automation, firm-wide deployment — is clear.")

add_divider(doc)

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run("CA Gopi  |  ZenCFO, Chennai  |  zencfoit@gmail.com  |  ICAI AI Hackathon Season 5")
set_font(r_f, size=9, color=(130, 130, 130), italic=True)

# ── Save ──
out_path = r"D:\ICAI Hackathon\Certify Genie\Certify_Genie_Article.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
