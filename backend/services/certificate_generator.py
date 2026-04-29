"""
Certificate Generator Service
------------------------------
Generates professional ICAI-format certificates with CA letterhead in PDF and DOCX.
Supports: Net Worth Certificate, Turnover Certificate

Design rules:
- ICAI logo in letterhead (top-left), firm name beside it
- UDIN printed in bold before the signature line
- CA name in signature block larger than surrounding text
- Rs. prefix in PDF (Helvetica has no Unicode Rupee glyph); ₹ in DOCX / HTML
- Meaningful display filenames
- Zero product branding in output
"""
import io
import logging
import os
import re
import uuid
from datetime import datetime
from typing import Optional

from config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()

NAVY   = (30, 58, 95)
NAVY_H = "#1E3A5F"
GOLD_H = "#C9A227"
GREY_H = "#555555"

# Path to the ICAI logo (webp → converted to PNG in memory for reportlab/docx)
_LOGO_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "Assets", "New-CA-Logo_ICAI.webp",
)


# ─── Logo Loader ──────────────────────────────────────────────────────────────

def _logo_png_bytes() -> Optional[io.BytesIO]:
    """Load the ICAI webp logo and return it as PNG bytes (for reportlab & docx)."""
    try:
        from PIL import Image
        img = Image.open(_LOGO_PATH).convert("RGBA")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception as e:
        logger.warning("Could not load ICAI logo: %s", e)
        return None


def _logo_rl_image(width_pts: float):
    """Return a reportlab Image flowable for the logo, preserving aspect ratio."""
    from reportlab.platypus import Image as RLImage
    buf = _logo_png_bytes()
    if buf is None:
        return None
    # Original: 761 × 587 pixels
    aspect = 587 / 761
    height_pts = width_pts * aspect
    return RLImage(buf, width=width_pts, height=height_pts)


# ─── Number Formatters ────────────────────────────────────────────────────────

def _format_inr_pdf(value: Optional[float]) -> str:
    if value is None:
        return "N/A"
    parts = str(int(abs(value)))
    formatted = parts if len(parts) <= 3 else parts[-3:]
    remaining = parts[:-3] if len(parts) > 3 else ""
    while remaining:
        formatted = remaining[-2:] + "," + formatted
        remaining = remaining[:-2]
    return f"Rs.{'-' if value < 0 else ''}{formatted}"


def _format_inr_docx(value: Optional[float]) -> str:
    if value is None:
        return "N/A"
    parts = str(int(abs(value)))
    formatted = parts if len(parts) <= 3 else parts[-3:]
    remaining = parts[:-3] if len(parts) > 3 else ""
    while remaining:
        formatted = remaining[-2:] + "," + formatted
        remaining = remaining[:-2]
    return f"\u20b9{'-' if value < 0 else ''}{formatted}"


# ─── Misc Helpers ─────────────────────────────────────────────────────────────

def _cert_number() -> str:
    now = datetime.now()
    return f"CERT-{now.year}-{now.month:02d}-{uuid.uuid4().hex[:4].upper()}"


def _build_display_filename(cert_number, company_name, financial_year, certificate_type):
    labels = {
        "net_worth": "Net-Worth-Certificate",
        "turnover": "Turnover-Certificate",
        "working_capital": "Working-Capital-Certificate",
    }
    safe_co = "-".join(re.sub(r"[^\w\s]", "", company_name or "Client").strip().split()[:5])
    safe_fy = re.sub(r"[^\w\-]", "", (financial_year or "").replace("/", "-").replace(" ", ""))
    label   = labels.get(certificate_type, certificate_type.replace("_", "-").title())
    return "_".join(p for p in [cert_number, safe_co, safe_fy, label] if p)


# ─── PDF Styles ───────────────────────────────────────────────────────────────

def _make_styles():
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    navy = HexColor(NAVY_H)
    base = getSampleStyleSheet()

    def ps(name, **kw):
        return ParagraphStyle(name, parent=base["Normal"], **kw)

    return {
        "navy":     navy,
        "title":    ps("CertTitle",  fontName="Helvetica-Bold", fontSize=16,
                        textColor=navy, alignment=TA_CENTER, spaceBefore=0, spaceAfter=4),
        "subtitle": ps("CertSub",    fontName="Helvetica",      fontSize=10,
                        textColor=navy, alignment=TA_CENTER, spaceAfter=2),
        "label":    ps("Label",      fontName="Helvetica",      fontSize=8,
                        textColor=HexColor(GREY_H), alignment=TA_CENTER, spaceAfter=2),
        "body":     ps("Body",       fontName="Helvetica",      fontSize=10,
                        leading=14, spaceAfter=4),
        "small":    ps("Small",      fontName="Helvetica",      fontSize=8.5,
                        leading=12, spaceAfter=2, textColor=HexColor(GREY_H)),
        "footer":   ps("Footer",     fontName="Helvetica",      fontSize=7.5,
                        textColor=HexColor("#888888"), alignment=TA_CENTER),
        "firm_h":   ps("FirmH",      fontName="Helvetica-Bold", fontSize=15,
                        textColor=navy, leading=18, spaceAfter=2),
        "firm_sub": ps("FirmSub",    fontName="Helvetica",      fontSize=9.5,
                        textColor=HexColor("#222222"), leading=13, spaceAfter=2),
        "firm_det": ps("FirmDet",    fontName="Helvetica",      fontSize=8,
                        textColor=HexColor(GREY_H), leading=11, spaceAfter=1),
    }


# ─── Letterhead Block (PDF) ───────────────────────────────────────────────────

def _letterhead_pdf(ca_details: dict, styles: dict, content_width_pts: float):
    """
    Professional CA letterhead:
    Left cell  — ICAI logo (or fallback navy badge)
    Right cell — firm name + CA details
    """
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.colors import HexColor

    LOGO_W   = 60          # pts
    LOGO_H   = LOGO_W * (587 / 761)
    firm_w   = content_width_pts - LOGO_W - 10

    logo_img = _logo_rl_image(LOGO_W)

    # Fallback: navy text badge when logo missing
    if logo_img is None:
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
        fb_style = ParagraphStyle("FB", fontName="Helvetica-Bold", fontSize=20,
                                  textColor=HexColor("#FFFFFF"), alignment=TA_CENTER)
        logo_cell = [Spacer(1, 6), Paragraph("CA", fb_style), Spacer(1, 6)]
        logo_bg   = HexColor(NAVY_H)
    else:
        logo_cell = [Spacer(1, 4), logo_img, Spacer(1, 4)]
        logo_bg   = HexColor("#FFFFFF")

    firm_name = (ca_details.get("firm_name") or "CA Firm").upper()
    ca_name   =  ca_details.get("name", "")
    mem_no    =  ca_details.get("membership_no", "")
    address   =  ca_details.get("address", "")

    firm_cell = [
        Paragraph(firm_name, styles["firm_h"]),
        Paragraph("Chartered Accountants", styles["firm_sub"]),
        Paragraph(
            f"{ca_name} &nbsp;&nbsp;|&nbsp;&nbsp; Membership No: {mem_no}",
            styles["firm_det"],
        ),
        Paragraph(address, styles["firm_det"]),
    ]

    tbl = Table([[logo_cell, firm_cell]], colWidths=[LOGO_W + 10, firm_w])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (0, 0), logo_bg),
        ("ALIGN",         (0, 0), (0, 0), "CENTER"),
        ("VALIGN",        (0, 0), (0, 0), "MIDDLE"),
        ("LEFTPADDING",   (0, 0), (0, 0), 4),
        ("RIGHTPADDING",  (0, 0), (0, 0), 4),
        ("TOPPADDING",    (0, 0), (0, 0), 4),
        ("BOTTOMPADDING", (0, 0), (0, 0), 4),
        ("LINEAFTER",     (0, 0), (0, 0), 0.5, HexColor("#CCCCCC")),
        ("VALIGN",        (1, 0), (1, 0), "MIDDLE"),
        ("LEFTPADDING",   (1, 0), (1, 0), 14),
        ("TOPPADDING",    (1, 0), (1, 0), 5),
        ("BOTTOMPADDING", (1, 0), (1, 0), 5),
        ("BOX",           (0, 0), (-1, -1), 0.75, HexColor(NAVY_H)),
    ]))
    return tbl


def _header_rule(navy_hex):
    """Thick + thin double rule after letterhead."""
    from reportlab.platypus import Table, TableStyle, Spacer
    from reportlab.lib.colors import HexColor
    zero_pad = [("TOPPADDING", (0,0),(-1,-1), 0), ("BOTTOMPADDING", (0,0),(-1,-1), 0)]
    thick = Table([[""]], colWidths=["100%"])
    thick.setStyle(TableStyle([("LINEBELOW", (0, 0), (-1, -1), 2.5, navy_hex)] + zero_pad))
    thin  = Table([[""]], colWidths=["100%"])
    thin.setStyle(TableStyle([("LINEBELOW", (0, 0), (-1, -1), 0.5, navy_hex)] + zero_pad))
    return [thick, Spacer(1, 2), thin]


# ─── Signature Block (PDF) ────────────────────────────────────────────────────

def _sig_block_pdf(elements, ca_details: dict, cert_date: str,
                   cert_number: str, styles: dict, udin: str = ""):
    from reportlab.platypus import Spacer, Table, TableStyle, Paragraph
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT

    elements.append(Spacer(1, 8))

    navy = HexColor(NAVY_H)

    ca_name_style = ParagraphStyle(
        "CANameSig", fontName="Helvetica-Bold", fontSize=11,
        textColor=navy, leading=13,
    )
    udin_style = ParagraphStyle(
        "UDINStyle", fontName="Helvetica-Bold", fontSize=9,
        textColor=HexColor("#111111"), leading=12,
    )

    sig_rows = [
        [Paragraph(f"For {ca_details.get('firm_name', '')}", styles["body"]), ""],
        [Paragraph("Chartered Accountants", styles["small"]),                 ""],
        ["", ""],  # gap row — signature line drawn below this
        [Paragraph(ca_details.get("name", ""), ca_name_style),               ""],
        [Paragraph(f"Membership No: {ca_details.get('membership_no', '')}", styles["small"]), ""],
        [Paragraph(ca_details.get("address", ""), styles["small"]),
         Paragraph(f"Date: {cert_date}", styles["small"])],
    ]
    if udin and udin.strip():
        sig_rows.append([Paragraph(f"UDIN: {udin.strip()}", udin_style), ""])

    tbl = Table(sig_rows, colWidths=["65%", "35%"],
                rowHeights=[None, None, 18, None, None, None] + ([None] if udin and udin.strip() else []))
    tbl_style = [
        ("FONTSIZE",      (0, 0), (-1, -1), 9),
        ("TOPPADDING",    (0, 0), (-1, -1), 1),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
        ("ALIGN",         (1, 5), (1, 5), "RIGHT"),
        ("LINEBELOW",     (0, 2), (0, 2), 0.5, HexColor("#AAAAAA")),
    ]
    tbl.setStyle(TableStyle(tbl_style))
    elements.append(tbl)
    elements.append(Spacer(1, 8))

    sep = Table([[""]], colWidths=["100%"])
    sep.setStyle(TableStyle([("LINEABOVE", (0, 0), (-1, -1), 0.4, HexColor("#AAAAAA"))]))
    elements.append(sep)
    elements.append(Spacer(1, 3))
    elements.append(Paragraph(f"Certificate No: {cert_number}", styles["footer"]))


# ─── Shared PDF Header ────────────────────────────────────────────────────────

def _pdf_header(elements, ca_details, s, content_w, cert_title,
                cert_number, cert_date, company, fy, subtitle=None):
    """Letterhead + rules + certificate title block + client info row."""
    from reportlab.platypus import Spacer, Table, TableStyle, Paragraph
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT

    navy = s["navy"]

    # ── Letterhead + thick double rule ──
    elements.append(_letterhead_pdf(ca_details, s, content_w))
    elements.append(Spacer(1, 2))
    elements += _header_rule(navy)
    elements.append(Spacer(1, 4))

    # ── Certificate title (larger) ──
    elements.append(Paragraph(cert_title, s["title"]))
    if subtitle:
        elements.append(Paragraph(subtitle, s["subtitle"]))
    elements.append(Spacer(1, 3))

    # ── Cert number LEFT — Date RIGHT ──
    num_style = ParagraphStyle("CertNumL", fontName="Helvetica-Bold", fontSize=8.5,
                               textColor=HexColor("#333333"), alignment=TA_LEFT)
    date_style = ParagraphStyle("CertDateR", fontName="Helvetica", fontSize=8.5,
                                textColor=HexColor("#555555"), alignment=TA_RIGHT)
    num_date = Table(
        [[Paragraph(f"Certificate No: {cert_number}", num_style),
          Paragraph(f"Date: {cert_date}", date_style)]],
        colWidths=["60%", "40%"],
    )
    num_date.setStyle(TableStyle([
        ("TOPPADDING",    (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ("LEFTPADDING",   (0, 0), (0, -1), 0),
        ("RIGHTPADDING",  (1, 0), (1, -1), 0),
    ]))
    elements.append(num_date)
    elements.append(Spacer(1, 3))

    # ── Thin rule below number/date ──
    thin = Table([[""]], colWidths=["100%"])
    thin.setStyle(TableStyle([
        ("LINEBELOW",     (0, 0), (-1, -1), 0.5, navy),
        ("TOPPADDING",    (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    elements.append(thin)
    elements.append(Spacer(1, 5))

    # ── Client info (2 rows: company / FY) ──
    if company and fy:
        info = Table(
            [
                ["Client / Company:", company],
                ["Financial Year:",   fy],
            ],
            colWidths=["28%", "72%"],
        )
        info.setStyle(TableStyle([
            ("FONTNAME",      (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1, -1), 9),
            ("TOPPADDING",    (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("BACKGROUND",    (0, 0), (-1, -1), HexColor("#F4F6FB")),
            ("BOX",           (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
            ("LINEAFTER",     (0, 0), (0, -1), 0.5, HexColor("#CCCCCC")),
            ("LINEBELOW",     (0, 0), (-1, 0), 0.5, HexColor("#CCCCCC")),
        ]))
        elements.append(info)
        elements.append(Spacer(1, 7))


# ─── Method-aware row builder (PDF) ──────────────────────────────────────────

def _build_nw_pdf_rows(method_id: str, extracted: dict, labels: dict) -> list:
    """
    Build the table rows for the net worth PDF based on the selected method.
    Each row: [Particulars label, Formatted amount string]
    """
    def v(key):
        return (extracted.get(key) or {}).get("value")

    def fmt(val):
        return _format_inr_pdf(val)

    def paren(val):
        return f"({fmt(val)})" if val else "Nil"

    sc = v("share_capital");       rs = v("reserves_surplus")
    al = v("accumulated_losses");  ta = v("total_assets");  tl = v("total_liabilities")
    ia = v("intangible_assets");   gw = v("goodwill") or 0
    rr = v("revaluation_reserve") or 0;  de = v("deferred_expenditure") or 0
    nw = v("net_worth")

    bv  = (sc + (rs or 0) - (al or 0)) if sc is not None and rs is not None else None
    tnw = (bv - (ia or 0) - gw) if bv is not None else None

    if method_id == "nav":
        return [
            [labels.get("total_assets",       "Total Assets (A)"),           fmt(ta)],
            [labels.get("total_liabilities",  "Less: Total Liabilities (B)"), paren(tl)],
            [labels.get("net_worth",          "Net Worth / NAV (A \u2212 B)"), fmt(nw)],
        ]
    elif method_id == "tangible_nw":
        return [
            [labels.get("share_capital",      "Share Capital (A)"),                 fmt(sc)],
            [labels.get("reserves_surplus",   "Add: Reserves & Surplus (B)"),       fmt(rs)],
            [labels.get("accumulated_losses", "Less: Accumulated Losses (C)"),      paren(al)],
            ["Book Value Net Worth (D = A+B\u2212C)",                               fmt(bv)],
            [labels.get("intangible_assets",  "Less: Intangible Assets (E)"),       paren(ia)],
            [labels.get("goodwill",           "Less: Goodwill (F)"),                paren(gw) if gw else "Nil"],
            [labels.get("net_worth",          "Tangible Net Worth (D \u2212 E \u2212 F)"), fmt(nw)],
        ]
    elif method_id == "adjusted_tbv":
        return [
            [labels.get("share_capital",       "Share Capital (A)"),                   fmt(sc)],
            [labels.get("reserves_surplus",    "Add: Reserves & Surplus (B)"),         fmt(rs)],
            [labels.get("accumulated_losses",  "Less: Accumulated Losses (C)"),        paren(al)],
            ["Book Value Net Worth (D = A+B\u2212C)",                                  fmt(bv)],
            [labels.get("intangible_assets",   "Less: Intangible Assets (E)"),         paren(ia)],
            [labels.get("goodwill",            "Less: Goodwill (F)"),                  paren(gw) if gw else "Nil"],
            ["Tangible Net Worth (G = D\u2212E\u2212F)",                               fmt(tnw)],
            [labels.get("revaluation_reserve", "Less: Revaluation Reserve (H)"),       paren(rr) if rr else "Nil"],
            [labels.get("net_worth",           "Adjusted TBV (G \u2212 H)"),           fmt(nw)],
        ]
    elif method_id == "companies_act":
        return [
            [labels.get("share_capital",        "Paid-up Share Capital (A)"),              fmt(sc)],
            [labels.get("reserves_surplus",     "Add: Free Reserves & Surplus (B)"),       fmt(rs)],
            [labels.get("accumulated_losses",   "Less: Accumulated Losses / Deficit (C)"), paren(al)],
            [labels.get("deferred_expenditure", "Less: Deferred Expenditure (D)"),         paren(de) if de else "Nil"],
            [labels.get("net_worth",            "Net Worth [Sec. 2(57)] (A+B\u2212C\u2212D)"), fmt(nw)],
        ]
    else:  # book_value, sebi_lodr, or empty
        return [
            [labels.get("share_capital",      "Share Capital (A)"),             fmt(sc)],
            [labels.get("reserves_surplus",   "Add: Reserves & Surplus (B)"),   fmt(rs)],
            [labels.get("accumulated_losses", "Less: Accumulated Losses (C)"),
             f"({fmt(al)})" if al else "Nil"],
            [labels.get("net_worth",          "Net Worth (A + B \u2212 C)"),    fmt(nw)],
        ]


# ─── Year-over-Year PDF Table ────────────────────────────────────────────────

def _build_yoy_pdf_table(extracted: dict, cert_type: str, content_w: float):
    """Return a YoY comparison Table flowable, or None if no previous_year data."""
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER

    py = extracted.get("previous_year")
    if not py or not isinstance(py, dict):
        return None
    if not any(py.get(k) is not None for k in ("net_worth", "turnover", "total_assets")):
        return None

    prev_fy = py.get("financial_year") or "Previous Year"
    curr_fy = extracted.get("financial_year") or "Current Year"

    def cv(key):
        return (extracted.get(key) or {}).get("value")

    def pv(key):
        val = py.get(key)
        return float(val) if val is not None else None

    def fmt(val):
        return _format_inr_pdf(val) if val is not None else "N/A"

    def pct(curr, prev):
        if curr is None or prev is None or prev == 0:
            return "—"
        change = ((curr - prev) / abs(prev)) * 100
        sign = "+" if change >= 0 else ""
        return f"{sign}{change:.1f}%"

    if cert_type == "net_worth":
        field_rows = [
            ("Share Capital",      "share_capital"),
            ("Reserves & Surplus", "reserves_surplus"),
            ("Net Worth",          "net_worth"),
            ("Total Assets",       "total_assets"),
            ("Total Liabilities",  "total_liabilities"),
        ]
    else:
        field_rows = [
            ("Turnover",      "turnover"),
            ("Net Profit",    "net_profit"),
            ("Total Assets",  "total_assets"),
        ]

    header = [["Particulars", prev_fy, curr_fy, "Change"]]
    rows = []
    for label, key in field_rows:
        c = cv(key)
        p = pv(key)
        rows.append([label, fmt(p), fmt(c), pct(c, p)])

    tbl_data = header + rows
    nw_col_w = content_w * 0.40
    other_col_w = (content_w - nw_col_w) / 3

    tbl = Table(tbl_data, colWidths=[nw_col_w, other_col_w, other_col_w, other_col_w])
    navy = HexColor("#1E3A5F")
    tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0),  navy),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  HexColor("#FFFFFF")),
        ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, -1), 8.5),
        ("ALIGN",         (1, 0), (-1, -1), "RIGHT"),
        ("ALIGN",         (0, 0), (0, -1),  "LEFT"),
        ("GRID",          (0, 0), (-1, -1), 0.3, HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1), [HexColor("#FFFFFF"), HexColor("#F8FAFC")]),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING",   (0, 0), (-1, -1), 8),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
    ]))
    return tbl


# ─── Net Worth PDF ────────────────────────────────────────────────────────────

def generate_net_worth_pdf(
    extracted, ca_details, cert_number, output_path,
    custom_narrative=None, additional_notes=None,
    custom_row_labels=None, udin="", nw_method_label="", nw_method_id="",
):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    PAGE_W, _ = A4
    M = 20 * mm
    content_w = PAGE_W - 2 * M

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            rightMargin=M, leftMargin=M,
                            topMargin=12*mm, bottomMargin=15*mm)
    s      = _make_styles()
    navy   = s["navy"]
    labels = custom_row_labels or {}

    company   = extracted.get("company_name") or "The Client"
    fy        = extracted.get("financial_year") or "FY 2024-25"
    cert_date = ca_details.get("date") or datetime.now().strftime("%d/%m/%Y")

    ta_val = (extracted.get("total_assets")      or {}).get("value")
    tl_val = (extracted.get("total_liabilities") or {}).get("value")
    nw_val = (extracted.get("net_worth")         or {}).get("value")

    elements = []
    _pdf_header(elements, ca_details, s, content_w,
                "NET WORTH CERTIFICATE", cert_number, cert_date, company, fy)

    narrative = custom_narrative or (
        f"This is to certify that based on the audited financial statements of "
        f"<b>{company}</b> for the financial year <b>{fy}</b>, "
        f"the Net Worth has been computed as under:"
    )
    elements.append(Paragraph(narrative, s["body"]))
    elements.append(Spacer(1, 4))

    method_rows = _build_nw_pdf_rows(nw_method_id, extracted, labels)
    tbl_data = [["Particulars", "Amount (Rs.)"]] + method_rows
    tbl = Table(tbl_data, colWidths=["65%", "35%"])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0),  navy),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  HexColor("#FFFFFF")),
        ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, -1), 9.5),
        ("ALIGN",         (1, 0), (1, -1),  "RIGHT"),
        ("FONTNAME",      (0, -1), (-1, -1), "Helvetica-Bold"),
        ("BACKGROUND",    (0, -1), (-1, -1), HexColor("#EEF4EE")),
        ("LINEABOVE",     (0, -1), (-1, -1), 1.5, navy),
        ("GRID",          (0, 0), (-1, -1),  0.3, HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS",(0, 1), (-1, -2), [HexColor("#FFFFFF"), HexColor("#F8FAFC")]),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING",   (0, 0), (-1, -1), 10),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 10),
    ]))
    elements.append(tbl)
    elements.append(Spacer(1, 5))

    if ta_val is not None and tl_val is not None:
        elements.append(Paragraph(
            f"Total Assets: {_format_inr_pdf(ta_val)}  |  Total Liabilities: {_format_inr_pdf(tl_val)}",
            s["small"],
        ))

    elements.append(Spacer(1, 4))
    elements.append(Paragraph(
        "The above Net Worth has been computed on the basis of the audited Balance Sheet "
        "as at the end of the above-mentioned financial year. This certificate is issued "
        "at the request of the client for the purpose mentioned by them.",
        s["body"],
    ))
    if nw_method_label:
        from reportlab.lib.styles import ParagraphStyle
        method_style = ParagraphStyle(
            "MethodAttr", fontName="Helvetica-Oblique", fontSize=8.5,
            textColor=HexColor("#555555"), leading=12, spaceAfter=4,
        )
        elements.append(Paragraph(
            f"Net Worth computed as per: <b>{nw_method_label}</b>",
            method_style,
        ))

    if additional_notes and additional_notes.strip():
        elements.append(Spacer(1, 6))
        for line in additional_notes.strip().split("\n"):
            if line.strip():
                elements.append(Paragraph(line.strip(), s["body"]))

    _sig_block_pdf(elements, ca_details, cert_date, cert_number, s, udin)
    doc.build(elements)
    logger.info("Net Worth PDF: %s", output_path)


# ─── Turnover PDF ─────────────────────────────────────────────────────────────

def generate_turnover_pdf(
    extracted, ca_details, cert_number, output_path,
    custom_narrative=None, additional_notes=None,
    custom_row_labels=None, udin="",
):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    PAGE_W, _ = A4
    M = 20 * mm
    content_w = PAGE_W - 2 * M

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            rightMargin=M, leftMargin=M,
                            topMargin=12*mm, bottomMargin=15*mm)
    s      = _make_styles()
    navy   = s["navy"]
    labels = custom_row_labels or {}

    company      = extracted.get("company_name") or "The Client"
    fy           = extracted.get("financial_year") or "FY 2024-25"
    cert_date    = ca_details.get("date") or datetime.now().strftime("%d/%m/%Y")
    turnover_val = (extracted.get("turnover")   or {}).get("value")
    profit_val   = (extracted.get("net_profit") or {}).get("value")

    elements = []
    _pdf_header(elements, ca_details, s, content_w,
                "TURNOVER CERTIFICATE", cert_number, cert_date, company, fy)

    narrative = custom_narrative or (
        f"This is to certify that the turnover of <b>{company}</b> "
        f"for the financial year <b>{fy}</b> is as under:"
    )
    elements.append(Paragraph(narrative, s["body"]))
    elements.append(Spacer(1, 8))

    tbl_data = [
        ["Particulars", "Amount (Rs.)"],
        [labels.get("turnover",   "Annual Turnover / Revenue from Operations"), _format_inr_pdf(turnover_val)],
        [labels.get("net_profit", "Net Profit / (Loss) After Tax"),             _format_inr_pdf(profit_val)],
    ]
    tbl = Table(tbl_data, colWidths=["65%", "35%"])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0),  navy),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  HexColor("#FFFFFF")),
        ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, -1), 10),
        ("ALIGN",         (1, 0), (1, -1),  "RIGHT"),
        ("GRID",          (0, 0), (-1, -1), 0.3, HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1), [HexColor("#FFFFFF"), HexColor("#F8FAFC")]),
        ("TOPPADDING",    (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING",   (0, 0), (-1, -1), 10),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 10),
    ]))
    elements.append(tbl)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph(
        "The above turnover figure is based on the audited financial statements of the company. "
        "This certificate is issued at the specific request of the client.",
        s["body"],
    ))

    if additional_notes and additional_notes.strip():
        elements.append(Spacer(1, 6))
        for line in additional_notes.strip().split("\n"):
            if line.strip():
                elements.append(Paragraph(line.strip(), s["body"]))

    _sig_block_pdf(elements, ca_details, cert_date, cert_number, s, udin)
    doc.build(elements)
    logger.info("Turnover PDF: %s", output_path)


# ─── Working Capital PDF ──────────────────────────────────────────────────────

def _build_wc_pdf_rows(extracted: dict, labels: dict) -> list:
    def v(key):
        return (extracted.get(key) or {}).get("value")
    def fmt(val):
        return _format_inr_pdf(val)
    ca = v("current_assets")
    cl = v("current_liabilities")
    wc = v("working_capital") if v("working_capital") is not None else (
        (ca - cl) if ca is not None and cl is not None else None
    )
    return [
        [labels.get("current_assets",      "Current Assets (A)"),             fmt(ca)],
        [labels.get("current_liabilities", "Less: Current Liabilities (B)"),  f"({fmt(cl)})" if cl else "Nil"],
        [labels.get("working_capital",     "Working Capital (A \u2212 B)"),   fmt(wc)],
    ]


def generate_working_capital_pdf(
    extracted, ca_details, cert_number, output_path,
    custom_narrative=None, additional_notes=None,
    custom_row_labels=None, udin="", wc_purpose="", **kwargs
):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    PAGE_W, _ = A4
    M = 20 * mm
    content_w = PAGE_W - 2 * M

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            rightMargin=M, leftMargin=M,
                            topMargin=12*mm, bottomMargin=15*mm)
    s      = _make_styles()
    navy   = s["navy"]
    labels = custom_row_labels or {}

    company   = extracted.get("company_name") or "The Client"
    fy        = extracted.get("financial_year") or "FY 2024-25"
    cert_date = ca_details.get("date") or datetime.now().strftime("%d/%m/%Y")

    elements = []
    _pdf_header(elements, ca_details, s, content_w,
                "WORKING CAPITAL CERTIFICATE", cert_number, cert_date, company, fy)

    narrative = custom_narrative or (
        f"This is to certify that based on the audited financial statements of "
        f"<b>{company}</b> for the financial year <b>{fy}</b>, "
        f"the Working Capital has been computed as under:"
    )
    elements.append(Paragraph(narrative, s["body"]))
    elements.append(Spacer(1, 4))

    wc_rows = _build_wc_pdf_rows(extracted, labels)
    tbl_data = [["Particulars", "Amount (Rs.)"]] + wc_rows
    tbl = Table(tbl_data, colWidths=["65%", "35%"])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0),  navy),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  HexColor("#FFFFFF")),
        ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, -1), 10),
        ("ALIGN",         (1, 0), (1, -1),  "RIGHT"),
        ("FONTNAME",      (0, -1), (-1, -1), "Helvetica-Bold"),
        ("BACKGROUND",    (0, -1), (-1, -1), HexColor("#EEF4EE")),
        ("LINEABOVE",     (0, -1), (-1, -1), 1.5, navy),
        ("GRID",          (0, 0), (-1, -1),  0.3, HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS",(0, 1), (-1, -2), [HexColor("#FFFFFF"), HexColor("#F8FAFC")]),
        ("TOPPADDING",    (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING",   (0, 0), (-1, -1), 10),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 10),
    ]))
    elements.append(tbl)
    elements.append(Spacer(1, 8))

    if wc_purpose and wc_purpose.strip():
        closing = (
            "The above Working Capital has been computed on the basis of the audited Balance Sheet "
            "as at the end of the above-mentioned financial year. This certificate is issued "
            f"at the specific request of the client for the purpose of: <b>{wc_purpose.strip()}</b>."
        )
    else:
        closing = (
            "The above Working Capital has been computed on the basis of the audited Balance Sheet "
            "as at the end of the above-mentioned financial year. This certificate is issued "
            "at the specific request of the client."
        )
    elements.append(Paragraph(closing, s["body"]))

    if additional_notes and additional_notes.strip():
        elements.append(Spacer(1, 6))
        for line in additional_notes.strip().split("\n"):
            if line.strip():
                elements.append(Paragraph(line.strip(), s["body"]))

    _sig_block_pdf(elements, ca_details, cert_date, cert_number, s, udin)
    doc.build(elements)
    logger.info("Working Capital PDF: %s", output_path)


# ─── DOCX Generator ──────────────────────────────────────────────────────────

def _build_nw_docx_rows(method_id: str, extracted: dict, labels: dict) -> list:
    """Build net worth table rows for DOCX based on selected method."""
    def v(key):
        return (extracted.get(key) or {}).get("value")

    def fmt(val):
        return _format_inr_docx(val)

    def paren(val):
        return f"({fmt(val)})" if val else "Nil"

    sc = v("share_capital");       rs = v("reserves_surplus")
    al = v("accumulated_losses");  ta = v("total_assets");  tl = v("total_liabilities")
    ia = v("intangible_assets");   gw = v("goodwill") or 0
    rr = v("revaluation_reserve") or 0;  de = v("deferred_expenditure") or 0
    nw = v("net_worth")

    bv  = (sc + (rs or 0) - (al or 0)) if sc is not None and rs is not None else None
    tnw = (bv - (ia or 0) - gw) if bv is not None else None

    if method_id == "nav":
        return [
            (labels.get("total_assets",       "Total Assets (A)"),           fmt(ta)),
            (labels.get("total_liabilities",  "Less: Total Liabilities (B)"), paren(tl)),
            (labels.get("net_worth",          "Net Worth / NAV (A \u2212 B)"), fmt(nw)),
        ]
    elif method_id == "tangible_nw":
        return [
            (labels.get("share_capital",      "Share Capital (A)"),                 fmt(sc)),
            (labels.get("reserves_surplus",   "Add: Reserves & Surplus (B)"),       fmt(rs)),
            (labels.get("accumulated_losses", "Less: Accumulated Losses (C)"),      paren(al)),
            ("Book Value Net Worth (D = A+B\u2212C)",                               fmt(bv)),
            (labels.get("intangible_assets",  "Less: Intangible Assets (E)"),       paren(ia)),
            (labels.get("goodwill",           "Less: Goodwill (F)"),                paren(gw) if gw else "Nil"),
            (labels.get("net_worth",          "Tangible Net Worth (D \u2212 E \u2212 F)"), fmt(nw)),
        ]
    elif method_id == "adjusted_tbv":
        return [
            (labels.get("share_capital",       "Share Capital (A)"),                   fmt(sc)),
            (labels.get("reserves_surplus",    "Add: Reserves & Surplus (B)"),         fmt(rs)),
            (labels.get("accumulated_losses",  "Less: Accumulated Losses (C)"),        paren(al)),
            ("Book Value Net Worth (D = A+B\u2212C)",                                  fmt(bv)),
            (labels.get("intangible_assets",   "Less: Intangible Assets (E)"),         paren(ia)),
            (labels.get("goodwill",            "Less: Goodwill (F)"),                  paren(gw) if gw else "Nil"),
            ("Tangible Net Worth (G = D\u2212E\u2212F)",                               fmt(tnw)),
            (labels.get("revaluation_reserve", "Less: Revaluation Reserve (H)"),       paren(rr) if rr else "Nil"),
            (labels.get("net_worth",           "Adjusted TBV (G \u2212 H)"),           fmt(nw)),
        ]
    elif method_id == "companies_act":
        return [
            (labels.get("share_capital",        "Paid-up Share Capital (A)"),              fmt(sc)),
            (labels.get("reserves_surplus",     "Add: Free Reserves & Surplus (B)"),       fmt(rs)),
            (labels.get("accumulated_losses",   "Less: Accumulated Losses / Deficit (C)"), paren(al)),
            (labels.get("deferred_expenditure", "Less: Deferred Expenditure (D)"),         paren(de) if de else "Nil"),
            (labels.get("net_worth",            "Net Worth [Sec. 2(57)] (A+B\u2212C\u2212D)"), fmt(nw)),
        ]
    else:
        al_val = v("accumulated_losses")
        return [
            (labels.get("share_capital",      "Share Capital (A)"),              fmt(sc)),
            (labels.get("reserves_surplus",   "Add: Reserves & Surplus (B)"),   fmt(rs)),
            (labels.get("accumulated_losses", "Less: Accumulated Losses (C)"),
             f"({fmt(al_val)})" if al_val else "Nil"),
            (labels.get("net_worth",          "Net Worth (A + B \u2212 C)"),    fmt(nw)),
        ]


def _build_wc_docx_rows(extracted: dict, labels: dict) -> list:
    def v(key):
        return (extracted.get(key) or {}).get("value")
    ca = v("current_assets")
    cl = v("current_liabilities")
    wc = v("working_capital") if v("working_capital") is not None else (
        (ca - cl) if ca is not None and cl is not None else None
    )
    return [
        (labels.get("current_assets",      "Current Assets (A)"),             _format_inr_docx(ca)),
        (labels.get("current_liabilities", "Less: Current Liabilities (B)"),  f"({_format_inr_docx(cl)})" if cl else "Nil"),
        (labels.get("working_capital",     "Working Capital (A \u2212 B)"),   _format_inr_docx(wc)),
    ]


def generate_docx(
    certificate_type, extracted, ca_details, cert_number, output_path,
    custom_narrative=None, additional_notes=None,
    custom_row_labels=None, udin="", nw_method_label="", nw_method_id="",
    wc_purpose="", **kwargs
):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    CONTENT_W = 16.0  # cm - A4 minus 2x2.5 cm margins
    ALT_ROWS  = ["FFFFFF", "F8FAFC"]

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    company   = extracted.get("company_name") or "The Client"
    fy        = extracted.get("financial_year") or "FY 2024-25"
    cert_date = ca_details.get("date") or datetime.now().strftime("%d/%m/%Y")
    labels    = custom_row_labels or {}

    # ── XML helpers ──────────────────────────────────────────────────────────

    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.lstrip("#"))
        tcPr.append(shd)

    def set_cell_margins(cell, top=40, start=80, bottom=40, end=80):
        tc    = cell._tc
        tcPr  = tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for side, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")
            tcMar.append(el)
        tcPr.append(tcMar)

    def set_table_border(table, color="CCCCCC", sz="4",
                         sides=("top", "left", "bottom", "right", "insideH", "insideV")):
        tbl    = table._tbl
        tblPr  = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblBdr = OxmlElement("w:tblBdr")
        for side in sides:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color.lstrip("#"))
            tblBdr.append(el)
        tblPr.append(tblBdr)

    def spacer():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        return p

    def styled_run(para, text, bold=False, italic=False, size=10, color=None, align=None):
        if align is not None:
            para.alignment = align
        r = para.add_run(text)
        r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor(*color)
        return r

    # ── Letterhead ───────────────────────────────────────────────────────────
    lh = doc.add_table(rows=1, cols=2)
    lh.autofit = False
    lh.columns[0].width = Cm(3.0)
    lh.columns[1].width = Cm(CONTENT_W - 3.0)

    logo_cell = lh.cell(0, 0)
    firm_cell = lh.cell(0, 1)

    logo_buf = _logo_png_bytes()
    logo_cell.paragraphs[0].clear()
    lp = logo_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_buf:
        lp.add_run().add_picture(logo_buf, width=Cm(2.6))
    else:
        set_cell_bg(logo_cell, NAVY_H)
        lr = lp.add_run("CA")
        lr.bold = True; lr.font.size = Pt(28); lr.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_margins(logo_cell, top=60, start=40, bottom=60, end=40)

    firm_cell.paragraphs[0].clear()
    fp = firm_cell.paragraphs[0]
    fp.paragraph_format.space_before = Pt(2)
    fr = fp.add_run((ca_details.get("firm_name") or "CA Firm").upper())
    fr.bold = True; fr.font.size = Pt(15); fr.font.color.rgb = RGBColor(*NAVY)

    def firm_line(text, size=9, rgb=(80, 80, 80)):
        p = firm_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.color.rgb = RGBColor(*rgb)

    firm_line("Chartered Accountants", size=9.5)
    firm_line(
        f"{ca_details.get('name', '')}  |  Membership No: {ca_details.get('membership_no', '')}",
        size=8.5,
    )
    firm_line(ca_details.get("address", ""), size=8)
    set_cell_margins(firm_cell, top=50, start=100, bottom=50, end=80)
    set_table_border(lh, color=NAVY_H, sz="6")

    # ── Double rule ──────────────────────────────────────────────────────────
    spacer()
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(0)
    rule_p.paragraph_format.space_after  = Pt(0)
    pPr  = rule_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "double")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), NAVY_H.lstrip("#"))
    pBdr.append(bot)
    pPr.append(pBdr)
    spacer()

    # ── Certificate title ────────────────────────────────────────────────────
    TITLES = {
        "net_worth":       "NET WORTH CERTIFICATE",
        "turnover":        "TURNOVER CERTIFICATE",
        "working_capital": "WORKING CAPITAL CERTIFICATE",
    }
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    tr = title_p.add_run(TITLES.get(certificate_type, "CERTIFICATE"))
    tr.bold = True; tr.font.size = Pt(16); tr.font.color.rgb = RGBColor(*NAVY)

    # ── Cert number (left) | Date (right) ────────────────────────────────────
    nd = doc.add_table(rows=1, cols=2)
    nd.autofit = False
    nd.columns[0].width = Cm(CONTENT_W * 0.60)
    nd.columns[1].width = Cm(CONTENT_W * 0.40)
    nd_l = nd.cell(0, 0); nd_r = nd.cell(0, 1)
    nd_l.paragraphs[0].clear()
    styled_run(nd_l.paragraphs[0], f"Certificate No: {cert_number}", bold=True, size=8.5)
    nd_r.paragraphs[0].clear()
    styled_run(nd_r.paragraphs[0], f"Date: {cert_date}", size=8.5,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    spacer()

    # ── Client info table (shaded 2-row — matches PDF) ────────────────────────
    ci = doc.add_table(rows=2, cols=2)
    ci.autofit = False
    ci.columns[0].width = Cm(CONTENT_W * 0.28)
    ci.columns[1].width = Cm(CONTENT_W * 0.72)
    for r_idx, (lbl_txt, val_txt) in enumerate(
        [("Client / Company:", company), ("Financial Year:", fy)]
    ):
        lc = ci.cell(r_idx, 0); vc = ci.cell(r_idx, 1)
        for cell in (lc, vc):
            set_cell_bg(cell, "F4F6FB")
            set_cell_margins(cell, top=40, start=80, bottom=40, end=80)
        lc.paragraphs[0].clear()
        styled_run(lc.paragraphs[0], lbl_txt, bold=True, size=9)
        vc.paragraphs[0].clear()
        styled_run(vc.paragraphs[0], val_txt, size=9)
    set_table_border(ci, color="CCCCCC", sz="4")
    spacer()

    # ── Narrative ────────────────────────────────────────────────────────────
    defaults = {
        "net_worth": (
            f"This is to certify that based on the audited financial statements of "
            f"{company} for the financial year {fy}, the Net Worth has been computed as under:"
        ),
        "turnover": (
            f"This is to certify that the turnover of {company} "
            f"for the financial year {fy} is as under:"
        ),
        "working_capital": (
            f"This is to certify that based on the audited financial statements of "
            f"{company} for the financial year {fy}, the Working Capital has been computed as under:"
        ),
    }
    narr_p = doc.add_paragraph(custom_narrative or defaults.get(certificate_type, ""))
    narr_p.paragraph_format.space_after = Pt(4)

    # ── Figures table ────────────────────────────────────────────────────────
    if certificate_type == "net_worth":
        rows_data = _build_nw_docx_rows(nw_method_id, extracted, labels)
    elif certificate_type == "turnover":
        tv = (extracted.get("turnover")   or {}).get("value")
        pv = (extracted.get("net_profit") or {}).get("value")
        rows_data = [
            (
                labels.get("turnover",   "Annual Turnover / Revenue from Operations"),
                _format_inr_docx(tv),
            ),
            (
                labels.get("net_profit", "Net Profit / (Loss) After Tax"),
                _format_inr_docx(pv),
            ),
        ]
    else:
        rows_data = _build_wc_docx_rows(extracted, labels)

    fig = doc.add_table(rows=len(rows_data) + 1, cols=2)
    fig.autofit = False
    fig.columns[0].width = Cm(CONTENT_W * 0.65)
    fig.columns[1].width = Cm(CONTENT_W * 0.35)

    # Header row
    for cell, txt in zip(fig.rows[0].cells, ["Particulars", "Amount (₹)"]):
        set_cell_bg(cell, NAVY_H)
        set_cell_margins(cell, top=50, start=80, bottom=50, end=80)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(txt)
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(255, 255, 255)
    fig.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Data rows — alternating backgrounds, last row green + bold
    for i, (lbl, amt) in enumerate(rows_data, 1):
        is_last = (i == len(rows_data))
        bg = "EEF4EE" if is_last else ALT_ROWS[(i - 1) % 2]
        lc = fig.rows[i].cells[0]; ac = fig.rows[i].cells[1]
        for cell in (lc, ac):
            set_cell_bg(cell, bg)
            set_cell_margins(cell, top=50, start=80, bottom=50, end=80)
        lc.paragraphs[0].clear()
        lr = lc.paragraphs[0].add_run(lbl)
        lr.bold = is_last; lr.font.size = Pt(9.5)
        ac.paragraphs[0].clear()
        ac.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ar = ac.paragraphs[0].add_run(amt)
        ar.bold = is_last; ar.font.size = Pt(9.5)

    set_table_border(fig, color="CCCCCC", sz="3")
    spacer()

    # ── Closing paragraph ─────────────────────────────────────────────────────
    wc_closing = (
        f"The above Working Capital has been computed on the basis of the audited Balance Sheet "
        f"as at the end of the above-mentioned financial year. This certificate is issued "
        f"at the specific request of the client for the purpose of: {wc_purpose.strip()}."
        if wc_purpose and wc_purpose.strip()
        else (
            "The above Working Capital has been computed on the basis of the audited Balance Sheet "
            "as at the end of the above-mentioned financial year. This certificate is issued "
            "at the specific request of the client."
        )
    )
    closing = {
        "net_worth": (
            "The above Net Worth has been computed on the basis of the audited Balance Sheet "
            "as at the end of the above-mentioned financial year. This certificate is issued "
            "at the request of the client for the purpose mentioned by them."
        ),
        "turnover": (
            "The above turnover figure is based on the audited financial statements of the company. "
            "This certificate is issued at the specific request of the client."
        ),
        "working_capital": wc_closing,
    }
    close_p = doc.add_paragraph(closing.get(certificate_type, ""))
    close_p.paragraph_format.space_after = Pt(2)

    if certificate_type == "net_worth" and nw_method_label:
        mp = doc.add_paragraph()
        mr = mp.add_run(f"Net Worth computed as per: {nw_method_label}")
        mr.italic = True; mr.font.size = Pt(8.5); mr.font.color.rgb = RGBColor(85, 85, 85)

    if additional_notes and additional_notes.strip():
        spacer()
        for line in additional_notes.strip().split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    # ── Signature block (table — mirrors PDF layout) ──────────────────────────
    spacer()
    sig = doc.add_table(rows=6, cols=2)
    sig.autofit = False
    sig.columns[0].width = Cm(CONTENT_W * 0.65)
    sig.columns[1].width = Cm(CONTENT_W * 0.35)

    def sig_para(row, col, text, bold=False, size=9, color=(60, 60, 60),
                 align=WD_ALIGN_PARAGRAPH.LEFT):
        cell = sig.cell(row, col)
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = align
        r = cell.paragraphs[0].add_run(text)
        r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = RGBColor(*color)
        set_cell_margins(cell, top=15, start=0, bottom=15, end=0)

    sig_para(0, 0, f"For {ca_details.get('firm_name', '')}", bold=True, size=10, color=(30, 30, 30))
    sig_para(1, 0, "Chartered Accountants", size=9, color=(80, 80, 80))

    # Row 2: signature gap with underline border
    sig.rows[2].height = Cm(1.2)
    sig_gap  = sig.cell(2, 0)
    set_cell_margins(sig_gap, top=0, start=0, bottom=0, end=0)
    tc    = sig_gap._tc
    tcPr  = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr")
    bb    = OxmlElement("w:bottom")
    bb.set(qn("w:val"), "single"); bb.set(qn("w:sz"), "4")
    bb.set(qn("w:space"), "0");    bb.set(qn("w:color"), "AAAAAA")
    tcBdr.append(bb)
    tcPr.append(tcBdr)

    sig_para(3, 0, ca_details.get("name", ""), bold=True, size=11, color=NAVY)
    sig_para(4, 0, f"Membership No: {ca_details.get('membership_no', '')}", size=8.5, color=(80, 80, 80))
    sig_para(5, 0, ca_details.get("address", ""), size=8.5, color=(80, 80, 80))
    sig_para(5, 1, f"Date: {cert_date}", size=8.5, color=(80, 80, 80),
             align=WD_ALIGN_PARAGRAPH.RIGHT)

    if udin and udin.strip():
        spacer()
        up = doc.add_paragraph()
        ur = up.add_run(f"UDIN: {udin.strip()}")
        ur.bold = True; ur.font.size = Pt(9)

    # ── Footer separator + cert number ───────────────────────────────────────
    spacer()
    sep_p    = doc.add_paragraph()
    sep_pPr  = sep_p._p.get_or_add_pPr()
    sep_pBdr = OxmlElement("w:pBdr")
    sep_top  = OxmlElement("w:top")
    sep_top.set(qn("w:val"),   "single")
    sep_top.set(qn("w:sz"),    "4")
    sep_top.set(qn("w:space"), "1")
    sep_top.set(qn("w:color"), "AAAAAA")
    sep_pBdr.append(sep_top)
    sep_pPr.append(sep_pBdr)

    fn_p = doc.add_paragraph()
    fn_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fn_r = fn_p.add_run(f"Certificate No: {cert_number}")
    fn_r.font.size = Pt(7.5); fn_r.font.color.rgb = RGBColor(136, 136, 136)

    doc.save(output_path)
    logger.info("%s DOCX: %s", certificate_type, output_path)

# ─── Dispatcher ──────────────────────────────────────────────────────────────

def _apply_pdf_password(pdf_path: str, password: str) -> None:
    """Encrypt an existing PDF with an open password using pypdf."""
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(pdf_path)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(user_password=password, owner_password=password, use_128bit=True)
    with open(pdf_path, "wb") as f:
        writer.write(f)
    logger.info("PDF password protection applied to %s", pdf_path)


def generate_certificate(
    certificate_type, extracted, ca_details, file_id,
    custom_narrative=None, additional_notes=None, custom_row_labels=None,
    custom_cert_number=None, custom_company_name=None, custom_financial_year=None,
    udin=None, nw_method_label=None, nw_method_id=None, pdf_password=None,
    wc_purpose=None,
) -> dict:
    cert_id     = uuid.uuid4().hex
    cert_number = (
        custom_cert_number.strip()
        if custom_cert_number and custom_cert_number.strip()
        else _cert_number()
    )

    output_dir = settings.output_dir
    os.makedirs(output_dir, exist_ok=True)

    if custom_company_name:
        extracted["company_name"] = custom_company_name
    if custom_financial_year:
        extracted["financial_year"] = custom_financial_year

    display_base = _build_display_filename(
        cert_number, extracted.get("company_name"),
        extracted.get("financial_year"), certificate_type,
    )
    pdf_path  = os.path.join(output_dir, f"{cert_id}.pdf")
    docx_path = os.path.join(output_dir, f"{cert_id}.docx")

    extra = {
        "custom_narrative":  custom_narrative,
        "additional_notes":  additional_notes,
        "custom_row_labels": custom_row_labels or {},
        "udin":              udin or "",
        "nw_method_label":   nw_method_label or "",
        "nw_method_id":      nw_method_id or "",
        "wc_purpose":        wc_purpose or "",
    }

    if certificate_type == "net_worth":
        generate_net_worth_pdf(extracted, ca_details, cert_number, pdf_path, **extra)
    elif certificate_type == "turnover":
        generate_turnover_pdf(extracted, ca_details, cert_number, pdf_path, **extra)
    elif certificate_type == "working_capital":
        generate_working_capital_pdf(extracted, ca_details, cert_number, pdf_path, **extra)
    else:
        raise ValueError(f"Unknown certificate type: {certificate_type}")

    generate_docx(certificate_type, extracted, ca_details, cert_number, docx_path, **extra)

    if pdf_password and pdf_password.strip():
        _apply_pdf_password(pdf_path, pdf_password.strip())

    return {
        "certificate_id":        cert_id,
        "cert_number":           cert_number,
        "pdf_path":              pdf_path,
        "docx_path":             docx_path,
        "pdf_display_filename":  f"{display_base}.pdf",
        "docx_display_filename": f"{display_base}.docx",
    }
